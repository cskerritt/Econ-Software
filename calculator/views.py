from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from .models import EconomicAnalysis, Evaluee, HealthcareCategory, HealthcarePlan, HealthcareCost
from .serializers import EconomicAnalysisSerializer, EvalueeSerializer, HealthcareCategorySerializer, HealthcarePlanSerializer, HealthcareCostSerializer
from openpyxl import Workbook
from docx import Document
from django.http import HttpResponse
from decimal import Decimal, ROUND_HALF_UP
from django.shortcuts import get_object_or_404

class EvalueeViewSet(viewsets.ModelViewSet):
    queryset = Evaluee.objects.all()
    serializer_class = EvalueeSerializer

class EconomicAnalysisViewSet(viewsets.ModelViewSet):
    queryset = EconomicAnalysis.objects.all()
    serializer_class = EconomicAnalysisSerializer

    def create(self, request, *args, **kwargs):
        try:
            return super().create(request, *args, **kwargs)
        except Exception as e:
            return Response(
                {'detail': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )

    def _calculate_benefits_loss(self, base_earnings, benefits_rate):
        """Calculate benefits loss based on base earnings and benefits rate"""
        return base_earnings * (benefits_rate / 100)

    def _calculate_insurance_loss(self, base_amount, growth_rate, years):
        """Calculate insurance loss with growth over years"""
        total_loss = 0
        current_amount = base_amount
        for year in range(years):
            total_loss += current_amount
            current_amount *= (1 + growth_rate / 100)
        return total_loss

    @action(detail=True, methods=['get'])
    def calculate(self, request, pk=None):
        try:
            analysis = self.get_object()
            evaluee = analysis.evaluee

            # Calculate age at injury and current age
            injury_date = analysis.date_of_injury
            report_date = analysis.date_of_report
            birth_date = evaluee.date_of_birth

            age_at_injury = (injury_date - birth_date).days / 365.25
            current_age = (report_date - birth_date).days / 365.25

            # Calculate retirement date and date of death
            retirement_date = injury_date.replace(year=injury_date.year + int(analysis.worklife_expectancy))
            date_of_death = injury_date.replace(year=injury_date.year + int(analysis.life_expectancy))

            # Personal Info
            personal_info = {
                'first_name': evaluee.first_name,
                'last_name': evaluee.last_name,
                'date_of_birth': evaluee.date_of_birth,
                'date_of_injury': analysis.date_of_injury,
                'date_of_report': analysis.date_of_report,
                'age_at_injury': round(age_at_injury, 1),
                'current_age': round(current_age, 1),
                'worklife_expectancy': analysis.worklife_expectancy,
                'years_to_final_separation': analysis.years_to_final_separation,
                'life_expectancy': analysis.life_expectancy,
                'retirement_date': retirement_date,
                'date_of_death': date_of_death,
            }

            # Pre-Injury Data
            pre_injury_rows = []
            total_future_value_pre = 0
            total_benefits_pre = 0
            total_insurance_pre = 0

            for row in analysis.pre_injury_rows.all():
                # Format portion of year as percentage
                portion_percentage = Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)
                
                # Calculate gross earnings (grown wage * portion of year)
                gross_earnings = row.wage_base_years * row.portion_of_year
                
                # Apply adjustment factor
                adjusted_earnings = gross_earnings * analysis.adjustment_factor

                # Calculate benefits and insurance losses
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    1  # For pre-injury period
                )

                total_future_value_pre += adjusted_earnings
                total_benefits_pre += benefits_loss
                total_insurance_pre += insurance_loss

                pre_injury_rows.append({
                    'year': row.year,
                    'portion_of_year': f"{portion_percentage}%",
                    'age': row.age,
                    'wage_base_years': row.wage_base_years,
                    'gross_earnings': gross_earnings,
                    'adjusted_earnings': adjusted_earnings,
                    'benefits_loss': benefits_loss,
                    'insurance_loss': insurance_loss
                })

            # Post-Injury Data
            post_injury_rows = []
            total_future_value_post = 0
            total_present_value_post = 0
            total_benefits_post = 0
            total_insurance_post = 0

            for row in analysis.post_injury_rows.all():
                # Format portion of year as percentage
                portion_percentage = Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)
                
                # Calculate gross earnings
                gross_earnings = row.wage_base_years * row.portion_of_year
                
                # Apply adjustment factor
                adjusted_earnings = gross_earnings * analysis.adjustment_factor

                # Calculate benefits and insurance losses
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    row.year - analysis.date_of_report.year + 1
                )

                # Apply discounting if enabled
                if analysis.apply_discounting and analysis.discount_rate:
                    years_from_report = row.year - analysis.date_of_report.year
                    discount_factor = (1 + analysis.discount_rate) ** years_from_report
                    present_value = adjusted_earnings / discount_factor
                    benefits_present_value = benefits_loss / discount_factor
                    insurance_present_value = insurance_loss / discount_factor
                    total_present_value_post += present_value
                    total_benefits_post += benefits_present_value
                    total_insurance_post += insurance_present_value
                else:
                    total_future_value_post += adjusted_earnings
                    total_benefits_post += benefits_loss
                    total_insurance_post += insurance_loss

                post_injury_rows.append({
                    'year': row.year,
                    'portion_of_year': f"{portion_percentage}%",
                    'age': row.age,
                    'wage_base_years': row.wage_base_years,
                    'gross_earnings': gross_earnings,
                    'adjusted_earnings': adjusted_earnings,
                    'benefits_loss': benefits_loss,
                    'insurance_loss': insurance_loss
                })

            # Prepare response data
            response_data = {
                'personal_info': personal_info,
                'exhibit1': {
                    'title': 'Pre-Injury Earnings',
                    'description': 'Earnings from date of injury to date of report',
                    'growth_rate': analysis.growth_rate,
                    'adjustment_factor': analysis.adjustment_factor,
                    'data': {
                        'rows': pre_injury_rows,
                        'total_future_value': total_future_value_pre,
                        'total_benefits': total_benefits_pre,
                        'total_insurance': total_insurance_pre
                    }
                },
                'exhibit2': {
                    'title': 'Post-Injury Earnings',
                    'description': 'Earnings loss from date of report to retirement (including residual capacity)',
                    'growth_rate': analysis.growth_rate,
                    'adjustment_factor': analysis.adjustment_factor,
                    'data': {
                        'rows': post_injury_rows,
                        'total_future_value': total_future_value_post,
                        'total_present_value': total_present_value_post if analysis.apply_discounting else None,
                        'total_benefits': total_benefits_post,
                        'total_insurance': total_insurance_post
                    }
                }
            }

            return Response(response_data)
        except Exception as e:
            return Response(
                {'detail': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['get'])
    def export_excel(self, request, pk=None):
        try:
            analysis = self.get_object()
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename=analysis_{analysis.id}.xlsx'

            workbook = Workbook()
            
            # Personal Info Sheet
            ws = workbook.active
            ws.title = "Personal Info"
            
            # Add headers and data for personal info
            personal_info = [
                ["First Name", analysis.evaluee.first_name],
                ["Last Name", analysis.evaluee.last_name],
                ["Date of Birth", analysis.evaluee.date_of_birth],
                ["Date of Injury", analysis.date_of_injury],
                ["Date of Report", analysis.date_of_report],
                ["Age at Injury", (analysis.date_of_injury - analysis.evaluee.date_of_birth).days / 365.25],
                ["Worklife Expectancy", analysis.worklife_expectancy],
                ["Life Expectancy", analysis.life_expectancy],
            ]
            
            for row in personal_info:
                ws.append(row)

            # Pre-Injury Sheet
            ws_pre = workbook.create_sheet("Pre-Injury Earnings")
            headers = ["Year", "Portion of Year", "Age", "Wage Base", "Gross Earnings", "Adjusted Earnings", "Benefits Loss", "Insurance Loss"]
            ws_pre.append(headers)
            
            for row in analysis.pre_injury_rows.all():
                gross_earnings = row.wage_base_years * row.portion_of_year
                adjusted_earnings = gross_earnings * analysis.adjustment_factor
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    1  # For pre-injury period
                )
                ws_pre.append([
                    row.year,
                    f"{Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)}%",
                    row.age,
                    row.wage_base_years,
                    gross_earnings,
                    adjusted_earnings,
                    benefits_loss,
                    insurance_loss
                ])

            # Post-Injury Sheet
            ws_post = workbook.create_sheet("Post-Injury Earnings")
            ws_post.append(headers)
            
            for row in analysis.post_injury_rows.all():
                gross_earnings = row.wage_base_years * row.portion_of_year
                adjusted_earnings = gross_earnings * analysis.adjustment_factor
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    row.year - analysis.date_of_report.year + 1
                )
                ws_post.append([
                    row.year,
                    f"{Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)}%",
                    row.age,
                    row.wage_base_years,
                    gross_earnings,
                    adjusted_earnings,
                    benefits_loss,
                    insurance_loss
                ])

            workbook.save(response)
            return response
            
        except Exception as e:
            return Response(
                {'detail': f'Failed to export Excel: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['get'])
    def export_word(self, request, pk=None):
        try:
            analysis = self.get_object()
            
            # Create a new document
            doc = Document()
            doc.add_heading('Economic Analysis Report', 0)
            
            # Add personal information section
            doc.add_heading('Personal Information', level=1)
            doc.add_paragraph(f'Name: {analysis.evaluee.first_name} {analysis.evaluee.last_name}')
            doc.add_paragraph(f'Date of Birth: {analysis.evaluee.date_of_birth}')
            doc.add_paragraph(f'Date of Injury: {analysis.date_of_injury}')
            doc.add_paragraph(f'Date of Report: {analysis.date_of_report}')
            
            # Add pre-injury earnings section
            doc.add_heading('Pre-Injury Earnings', level=1)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            headers = ["Year", "Portion of Year", "Age", "Wage Base", "Gross Earnings", "Adjusted Earnings", "Benefits Loss", "Insurance Loss"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                
            for row in analysis.pre_injury_rows.all():
                gross_earnings = row.wage_base_years * row.portion_of_year
                adjusted_earnings = gross_earnings * analysis.adjustment_factor
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    1  # For pre-injury period
                )
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.year)
                row_cells[1].text = f"{Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)}%"
                row_cells[2].text = f"{row.age:.1f}"
                row_cells[3].text = f"${row.wage_base_years:,.2f}"
                row_cells[4].text = f"${gross_earnings:,.2f}"
                row_cells[5].text = f"${adjusted_earnings:,.2f}"
                row_cells[6].text = f"${benefits_loss:,.2f}"
                row_cells[7].text = f"${insurance_loss:,.2f}"
                
            # Add post-injury earnings section
            doc.add_heading('Post-Injury Earnings', level=1)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                
            for row in analysis.post_injury_rows.all():
                gross_earnings = row.wage_base_years * row.portion_of_year
                adjusted_earnings = gross_earnings * analysis.adjustment_factor
                benefits_loss = self._calculate_benefits_loss(gross_earnings, analysis.benefits_rate)
                insurance_loss = self._calculate_insurance_loss(
                    analysis.health_insurance_base,
                    analysis.growth_rate,
                    row.year - analysis.date_of_report.year + 1
                )
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.year)
                row_cells[1].text = f"{Decimal(row.portion_of_year * 100).quantize(Decimal('0.1'), rounding=ROUND_HALF_UP)}%"
                row_cells[2].text = f"{row.age:.1f}"
                row_cells[3].text = f"${row.wage_base_years:,.2f}"
                row_cells[4].text = f"${gross_earnings:,.2f}"
                row_cells[5].text = f"${adjusted_earnings:,.2f}"
                row_cells[6].text = f"${benefits_loss:,.2f}"
                row_cells[7].text = f"${insurance_loss:,.2f}"
            
            # Save the document
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=analysis_{analysis.id}.docx'
            doc.save(response)
            
            return response
            
        except Exception as e:
            return Response(
                {'detail': f'Failed to export Word: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['get'], url_path='excel')
    def excel(self, request, pk=None):
        """Generate Excel report for the analysis"""
        try:
            analysis = self.get_object()
            
            # Create workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Economic Analysis"
            
            # Add headers
            ws['A1'] = "Economic Analysis Report"
            ws['A2'] = f"Evaluee: {analysis.evaluee.first_name} {analysis.evaluee.last_name}"
            ws['A3'] = f"Date of Injury: {analysis.date_of_injury}"
            ws['A4'] = f"Date of Report: {analysis.date_of_report}"
            
            # Add analysis details
            ws['A6'] = "Pre-Injury Base Wage"
            ws['B6'] = analysis.pre_injury_base_wage
            ws['A7'] = "Post-Injury Base Wage"
            ws['B7'] = analysis.post_injury_base_wage
            ws['A8'] = "Growth Rate"
            ws['B8'] = f"{analysis.growth_rate * 100}%"
            ws['A9'] = "Adjustment Factor"
            ws['B9'] = analysis.adjustment_factor
            
            if analysis.include_health_insurance:
                ws['A11'] = "Health Insurance"
                ws['B11'] = analysis.health_insurance_base
                ws['A12'] = "Health Cost Inflation Rate"
                ws['B12'] = f"{analysis.health_cost_inflation_rate * 100}%"
            
            if analysis.include_pension:
                ws['A14'] = "Pension Information"
                ws['B14'] = analysis.pension_type
                if analysis.pension_type == 'defined_benefit':
                    ws['A15'] = "Final Average Salary"
                    ws['B15'] = analysis.final_average_salary
                    ws['A16'] = "Years of Service"
                    ws['B16'] = analysis.years_of_service
                    ws['A17'] = "Benefit Multiplier"
                    ws['B17'] = f"{analysis.benefit_multiplier * 100}%"
                else:
                    ws['A15'] = "Annual Contribution"
                    ws['B15'] = analysis.annual_contribution
                    ws['A16'] = "Expected Return Rate"
                    ws['B16'] = f"{analysis.expected_return_rate * 100}%"
            
            # Create response
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename=analysis_{pk}.xlsx'
            wb.save(response)
            return response
        except Exception as e:
            return Response(
                {'detail': f'Failed to export Excel: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['get'], url_path='word')
    def word(self, request, pk=None):
        """Generate Word report for the analysis"""
        try:
            analysis = self.get_object()
            
            # Create document
            doc = Document()
            doc.add_heading('Economic Analysis Report', 0)
            
            # Add evaluee information
            doc.add_paragraph(f"Evaluee: {analysis.evaluee.first_name} {analysis.evaluee.last_name}")
            doc.add_paragraph(f"Date of Injury: {analysis.date_of_injury}")
            doc.add_paragraph(f"Date of Report: {analysis.date_of_report}")
            
            # Add analysis details
            doc.add_heading('Analysis Details', level=1)
            doc.add_paragraph(f"Pre-Injury Base Wage: ${analysis.pre_injury_base_wage:,.2f}")
            doc.add_paragraph(f"Post-Injury Base Wage: ${analysis.post_injury_base_wage:,.2f}")
            doc.add_paragraph(f"Growth Rate: {analysis.growth_rate * 100}%")
            doc.add_paragraph(f"Adjustment Factor: {analysis.adjustment_factor}")
            
            if analysis.include_health_insurance:
                doc.add_heading('Health Insurance', level=1)
                doc.add_paragraph(f"Base Amount: ${analysis.health_insurance_base:,.2f}")
                doc.add_paragraph(f"Inflation Rate: {analysis.health_cost_inflation_rate * 100}%")
            
            if analysis.include_pension:
                doc.add_heading('Pension Information', level=1)
                doc.add_paragraph(f"Pension Type: {analysis.pension_type}")
                if analysis.pension_type == 'defined_benefit':
                    doc.add_paragraph(f"Final Average Salary: ${analysis.final_average_salary:,.2f}")
                    doc.add_paragraph(f"Years of Service: {analysis.years_of_service}")
                    doc.add_paragraph(f"Benefit Multiplier: {analysis.benefit_multiplier * 100}%")
                else:
                    doc.add_paragraph(f"Annual Contribution: ${analysis.annual_contribution:,.2f}")
                    doc.add_paragraph(f"Expected Return Rate: {analysis.expected_return_rate * 100}%")
            
            # Create response
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=analysis_{pk}.docx'
            doc.save(response)
            return response
        except Exception as e:
            return Response(
                {'detail': f'Failed to export Word: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

class HealthcareCategoryViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = HealthcareCategory.objects.all()
    serializer_class = HealthcareCategorySerializer

class HealthcarePlanViewSet(viewsets.ModelViewSet):
    serializer_class = HealthcarePlanSerializer

    def get_queryset(self):
        analysis_id = self.kwargs['analysis_pk']
        return HealthcarePlan.objects.filter(analysis_id=analysis_id)

    def perform_create(self, serializer):
        analysis_id = self.kwargs['analysis_pk']
        analysis = get_object_or_404(EconomicAnalysis, id=analysis_id)
        serializer.save(analysis=analysis)

    @action(detail=True, methods=['post'])
    def toggle(self, request, analysis_pk=None, pk=None):
        plan = self.get_object()
        plan.is_active = not plan.is_active
        plan.save()
        return Response(self.get_serializer(plan).data)

    @action(detail=False, methods=['post'])
    def calculate_costs(self, request, analysis_pk=None):
        analysis = get_object_or_404(EconomicAnalysis, id=analysis_pk)
        plans = self.get_queryset().filter(is_active=True)
        
        # Delete existing costs
        HealthcareCost.objects.filter(plan__in=plans).delete()
        
        # Calculate costs for each plan
        for plan in plans:
            category = plan.category
            base_cost = plan.base_cost
            growth_rate = category.growth_rate
            frequency = category.frequency_years
            
            # Get the year range from the analysis
            start_year = analysis.date_of_injury.year
            end_year = start_year + int(analysis.life_expectancy)
            
            # Calculate age at start
            age_at_start = (analysis.date_of_injury - analysis.evaluee.date_of_birth).days / 365.25
            
            costs = []
            for year in range(start_year, end_year + 1):
                years_from_start = year - start_year
                age = age_at_start + years_from_start
                
                # Check if this is a treatment year based on frequency
                if frequency == 1 or years_from_start % int(frequency) == 0:
                    # Calculate grown cost
                    cost = base_cost * (1 + growth_rate) ** years_from_start
                    
                    costs.append(HealthcareCost(
                        plan=plan,
                        year=year,
                        age=age,
                        cost=cost
                    ))
            
            # Bulk create the costs
            HealthcareCost.objects.bulk_create(costs)
        
        return Response({'status': 'Costs calculated successfully'})
